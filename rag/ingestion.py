from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from pypdf import PdfReader
from docx import Document as DocxDocument

from config.settings import Settings
from rag.chunking import build_chunker
from rag.types import DocumentChunk

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".json", ".csv"}


def _document_id(path: Path) -> str:
    """Stable document id derived from path."""
    return hashlib.sha1(str(path).encode()).hexdigest()[:16]


def _infer_document_type(path: Path) -> str:
    stem_lower = path.stem.lower()
    if "resume" in stem_lower or "cv" in stem_lower:
        return "resume"
    if "experience" in stem_lower:
        return "experience"
    if "education" in stem_lower:
        return "education"
    if "skill" in stem_lower:
        return "skills"
    if "project" in stem_lower:
        return "projects"
    if "achievement" in stem_lower:
        return "achievements"
    if "certification" in stem_lower:
        return "certifications"
    if "publication" in stem_lower:
        return "publications"
    if "about" in stem_lower:
        return "about"
    if "faq" in stem_lower:
        return "faq"
    return "general"


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _load_json(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        data = json.loads(raw)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return raw


def _load_csv(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(raw))
    rows = [", ".join(row) for row in reader]
    return "\n".join(rows)


def load_document(path: Path) -> str:
    """Load document text from a file by extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(path)
    if ext == ".docx":
        return _load_docx(path)
    if ext in (".txt", ".md", ".markdown"):
        return _load_text(path)
    if ext == ".json":
        return _load_json(path)
    if ext == ".csv":
        return _load_csv(path)
    raise ValueError(f"Unsupported file extension: {ext}")


def load_document_bytes(content: bytes, filename: str) -> str:
    """Load document from raw bytes (e.g. uploaded via Streamlit)."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if suffix in (".txt", ".md", ".markdown"):
        return content.decode("utf-8", errors="replace")
    if suffix == ".json":
        raw = content.decode("utf-8", errors="replace")
        try:
            data = json.loads(raw)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            return raw
    if suffix == ".csv":
        raw = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(raw))
        return "\n".join(", ".join(row) for row in reader)
    raise ValueError(f"Unsupported file type: {suffix}")


def build_chunks_from_file(path: Path, settings: Settings) -> list[DocumentChunk]:
    """Full ingestion pipeline for a single file: load → chunk → metadata."""
    text = load_document(path)
    return _build_chunks(text, path.name, settings)


def build_chunks_from_bytes(content: bytes, filename: str, settings: Settings) -> list[DocumentChunk]:
    """Full ingestion pipeline for uploaded bytes: load → chunk → metadata."""
    text = load_document_bytes(content, filename)
    return _build_chunks(text, filename, settings)


def _build_chunks(text: str, filename: str, settings: Settings) -> list[DocumentChunk]:
    if not text.strip():
        return []

    source_path = Path(filename)
    doc_id = _document_id(source_path)
    doc_type = _infer_document_type(source_path)

    base_metadata: dict[str, Any] = {
        "source": filename,
        "document_type": doc_type,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }

    chunker = build_chunker(settings.chunk_size, settings.chunk_overlap)
    chunks = chunker.chunk_document(text, base_metadata, doc_id)
    logger.info(
        "Ingested document",
        extra={
            "extra_data": {
                "filename": filename,
                "chunks": len(chunks),
                "document_type": doc_type,
            }
        },
    )
    return chunks


def ingest_data_directory(settings: Settings, vector_store: Any) -> dict[str, int]:
    """Scan the data directory and ingest all supported documents."""
    data_dir = settings.data_dir
    results: dict[str, int] = {}

    for path in sorted(data_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            chunks = build_chunks_from_file(path, settings)
            relative = str(path.relative_to(data_dir))
            # Rewrite source to relative path for nicer display
            for chunk in chunks:
                chunk.metadata["source"] = relative
            added = vector_store.upsert_chunks(chunks)
            results[relative] = added
        except Exception as exc:
            logger.warning("Failed to ingest %s: %s", path, exc)

    return results
