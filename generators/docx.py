from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from config.settings import Settings


def _safe_filename(text: str) -> str:
    slug = re.sub(r"[^a-z0-9_]+", "_", text.lower()).strip("_")
    return slug[:40] or "document"


def build_docx(content: str, doc_type: str, settings: Settings) -> Path:
    """Generate a simple DOCX from Markdown-like content and return the file path."""
    doc = Document()

    # Style heading
    title_para = doc.add_heading(doc_type, level=1)
    title_para.runs[0].font.color.rgb = RGBColor(0x1F, 0x23, 0x28)

    doc.add_paragraph("")  # spacer

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    settings.generated_dir.mkdir(parents=True, exist_ok=True)
    slug = _safe_filename(doc_type)
    filename = f"{slug}_{uuid.uuid4().hex[:8]}.docx"
    output_path = settings.generated_dir / filename
    doc.save(str(output_path))
    return output_path
