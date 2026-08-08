from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from config.settings import Settings


def _safe_filename(text: str) -> str:
    slug = re.sub(r"[^a-z0-9_]+", "_", text.lower()).strip("_")
    return slug[:40] or "resume"


def build_resume_docx(resume_data: dict[str, Any], role: str, settings: Settings) -> Path:
    """Generate a tailored resume DOCX from structured resume data."""
    doc = Document()

    # Name / title
    name = str(resume_data.get("name") or "Resume")
    doc.add_heading(name, level=1).runs[0].font.color.rgb = RGBColor(0x1F, 0x23, 0x28)
    doc.add_paragraph(f"Target Role: {role}").runs[0].font.bold = True

    # Summary
    summary = str(resume_data.get("summary") or "")
    if summary:
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

    # Experience
    experience: list[Any] = resume_data.get("experience") or []
    if experience:
        doc.add_heading("Experience", level=2)
        for job in experience:
            if isinstance(job, dict):
                title_line = f"{job.get('title', '')} — {job.get('company', '')} ({job.get('dates', '')})"
                doc.add_paragraph(title_line).runs[0].font.bold = True
                for bullet in job.get("bullet_points", []):
                    doc.add_paragraph(str(bullet), style="List Bullet")
            else:
                doc.add_paragraph(str(job), style="List Bullet")

    # Education
    education: list[Any] = resume_data.get("education") or []
    if education:
        doc.add_heading("Education", level=2)
        for edu in education:
            if isinstance(edu, dict):
                edu_line = f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('dates', '')})"
                doc.add_paragraph(edu_line)
            else:
                doc.add_paragraph(str(edu))

    # Skills
    skills: list[Any] = resume_data.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(str(s) for s in skills))

    # Projects
    projects: list[Any] = resume_data.get("projects") or []
    if projects:
        doc.add_heading("Projects", level=2)
        for proj in projects:
            if isinstance(proj, dict):
                proj_line = f"{proj.get('name', '')} — {proj.get('description', '')}"
                doc.add_paragraph(proj_line, style="List Bullet")
                techs = proj.get("technologies")
                if techs:
                    doc.add_paragraph(f"Technologies: {techs}")
            else:
                doc.add_paragraph(str(proj), style="List Bullet")

    # Certifications
    certs: list[Any] = resume_data.get("certifications") or []
    if certs:
        doc.add_heading("Certifications", level=2)
        for cert in certs:
            doc.add_paragraph(str(cert), style="List Bullet")

    settings.generated_dir.mkdir(parents=True, exist_ok=True)
    slug = re.sub(r"[^a-z0-9_]+", "_", role.lower()).strip("_") or "general"
    filename = f"resume_{slug}_{uuid.uuid4().hex[:8]}.docx"
    output_path = settings.generated_dir / filename
    doc.save(str(output_path))
    return output_path
